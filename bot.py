import telegram
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
from docx import Document
from telegram import Update
from telegram.ext import Application, MessageHandler, filters, CallbackContext
import psycopg2
import os
import asyncio

bot = telegram.Bot(token='7702684837:AAG_y38JgEQeg1JkJ7yFDhGY285VTXQGgic')
model = SentenceTransformer("all-MiniLM-L6-v2")
dimension = 384
index = faiss.IndexFlatL2(dimension)

conn = psycopg2.connect(
    dbname="document_db", user="postgres", password="123", host="localhost", port="8080"
)
cursor = conn.cursor()

vectors = []

def save_document(text, doc_type="general"):
    cursor.execute(
        "INSERT INTO documents (text, doc_type) VALUES (%s, %s) RETURNING id", (text, doc_type)
    )
    document_id = cursor.fetchone()[0]
    conn.commit()
    return document_id

def save_vector(document_id, vector):
    vector_array = np.array(vector).tolist()
    cursor.execute(
        "INSERT INTO vectors (document_id, vector) VALUES (%s, %s)", (document_id, vector_array)
    )
    conn.commit()

def search_vector(query, index):
    query_vector = model.encode(query).astype('float32')
    D, I = index.search(np.array([query_vector]), k=1)
    result_text = vectors[I[0][0]][0]
    return result_text

async def handle_document(update: Update, context: CallbackContext):
    file = await update.message.document.get_file()
    filepath = os.path.join("downloads", f"{file.file_id}.docx")
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    await file.download(custom_path=filepath)
    parsed_data = parse_docx(filepath)
    text = " ".join(parsed_data["lists"] + parsed_data["tables"])
    document_id = save_document(text)
    vectors_data = vectorize_text(parsed_data)
    save_vector(document_id, vectors_data[0][1])

    for text, vector in vectors_data:
        index.add(np.array([vector]))
        vectors.append((text, vector))

    await update.message.reply_text("Документ и вектор добавлены в базу!")

def parse_docx(filepath):
    document = Document(filepath)
    data = {"tables": [], "lists": []}
    for table in document.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])
        data["tables"].append(table_data)
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith('List'):
            data["lists"].append(paragraph.text.strip())
    return data

def vectorize_text(data):
    vectors_data = []
    for table in data["tables"]:
        for row in table:
            text = " ".join(row)
            vector = model.encode(text)
            vectors_data.append((text, vector))
    for item in data["lists"]:
        vector = model.encode(item)
        vectors_data.append((item, vector))
    return vectors_data

async def handle_message(update: Update, context: CallbackContext):
    user_message = update.message.text
    response = search_vector(user_message, index)
    await update.message.reply_text(response)

async def main():
    application = Application.builder().token('7702684837:AAG_y38JgEQeg1JkJ7yFDhGY285VTXQGgic').build()
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))

    document_handler = MessageHandler(
        filters.Document.MimeType("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        handle_document)

    application.add_handler(document_handler)
    await application.run_polling()

if __name__ == "__main__":
    asyncio.run(main())

